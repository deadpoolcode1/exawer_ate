#!/usr/bin/env python3
"""Build the client-facing milestone hand-over .docx.

Separate from `build_docs_docx.py`, which converts our internal markdown docs
via pandoc. This one is written directly with python-docx because a hand-over
is a different artifact: two pages, three tables, and every claim in it has a
matching file in the evidence folder.

Regenerate after changing any of the numbers it quotes — they are deliberately
inline rather than computed, so that a stale figure is a visible edit in the
diff rather than something the script silently recalculates from a run that no
longer matches what was shipped.

    python scripts/build_handover_docx.py

Output path is the milestone package on the Desktop; change OUT for a new
milestone.
"""
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = "/home/ilan/Desktop/Exaware_M2_handover_2026-08-13/M2_Handover.docx"
#: Keep in step with deliverables/M2/ — every claim below has an evidence file.
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x59, 0x59, 0x59)


def style(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(6)


def h(doc, text, size=13, space_before=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = ACCENT
    return p


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(htxt)
        r.bold = True
        r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
    return t


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, tuple):
            lead, rest = it
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest).font.size = Pt(10.5)
        else:
            p.add_run(it).font.size = Pt(10.5)


doc = Document()
style(doc)

# ── title ───────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Milestone 2 — Hand-over")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = ACCENT

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("Dirty Queue & Code Generation · SOW PQ4476E\n"
              "CodeValue → Exaware · 13 August 2026")
r.font.size = Pt(10)
r.font.color.rgb = MUTED

# ── acceptance ──────────────────────────────────────────────────────────
h(doc, "M2 deliverables — all four met", space_before=14)
table(doc,
      ["SOW M2 deliverable", "Result"],
      [["Code generation based on selected tests",
        "TC01 / TC02 / TC03, emitted only for what the dirty queue marks SELECTED"],
       ["Pattern matching implementation",
        "537 of 612 automatable plan rows (87.7%) mapped to typed executable steps"],
       ["Demo: extract requirements from sample docs",
        "133 requirements → 269 plan rows → 698 action rows, from the SFS, CLI doc and 2 RFCs"],
       ["Up to 3 integration-ready test plans",
        "Three suites that compile unmodified against cmp-infra-project and cmp-tests-project"]])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run("Also delivered: the dirty queue itself (the SOW lists it under M4), "
              "per-scenario device configuration, and a device-verification stage "
              "that did not exist when M2 was scoped.")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED

# ── verified ────────────────────────────────────────────────────────────
h(doc, "Verified on your hardware")
p = doc.add_paragraph()
r = p.add_run("SUT pc-3080 / exa-il01-uf-3080, software 8.7.0 LAB 22, "
              "application build feature/dev64_evpn_23Jul2026.")
r.font.size = Pt(10)
bullets(doc, [
    ("All three suites run end to end - ",
     "TC01, TC02 and TC03 each complete under JUnit + JSystem against the DUT: "
     "OK (1 test), exit 0, full bring-up and tear-down included - which is where the "
     "previous milestone stopped. Afterwards show evpn detail lists evi-1 with its "
     "three attachment circuits bound. Please read the next section before reading "
     "\"green\" as \"the scenarios pass\"."),
    ("Compiles against your framework - ",
     "953 sources -> 1454 classes, zero errors; the generated files pass "
     "javac --release 8 -Werror -Xlint:all with zero warnings."),
    ("bringUpParams.crt passes your own validator - ",
     "TemplateManager.validateAgainstTemplate returns true (bringUpParameters_C0_002), "
     "standalone and inside the live bring-up. Adding one // line inside a table makes "
     "the same validator reject it, so the check is not vacuous."),
    ("One .cfg, two different platforms - ",
     "the same generated files ran unchanged on pc-3021 (edgeCore) and pc-3080 "
     "(UfiSpace); the int1/int2/int3 placeholders resolved to x-eth 0/0/8, 0/0/18 "
     "and 0/0/26 from your SUT file."),
    ("Every command the suite issues exists on this build - ",
     "ate capture reports 0 unsupported."),
])

h(doc, "Two defects the run found, which matter more than the pass", size=11,
  space_before=10)
bullets(doc, [
    ("A vlan-based EVI will not bind a port - ",
     "the commit is rejected: \"interface x-eth 0/0/8 is not a sub-interface, but the "
     "EVPN service-type is vlan-based\". This is in neither the SFS nor the CLI doc. "
     "The generator now creates the attachment circuits as sub-interfaces first, using "
     "the same stanza your VPLS suite uses."),
    ("A rejected command could not fail the test - ",
     "three configuration commands were refused by the CLI and the run stayed green: "
     "nothing was staged, so the commit had nothing to do, so configAndValidate logged "
     "a warning. Generated configuration steps now assert acceptance themselves. A "
     "negative control - an out-of-range sub-interface - turns the same run red, so we "
     "know the assertion works."),
])

h(doc, "One number went down on purpose", size=11, space_before=10)
p = doc.add_paragraph()
r = p.add_run("Usable captured expectations are 2 of 11, where the last hand-over said 7. "
              "Five of those seven were the MAC table's legend with no MAC address in it - "
              "an assertion that passes on any device, working or broken, and could never "
              "detect a regression. ate capture now refuses them and says why. The suite "
              "asserts less and means more; 02_evidence/evidence_capture_pc3080.txt shows "
              "one of the five in full.")
r.font.size = Pt(10.5)
p = doc.add_paragraph()
r = p.add_run("The cause is not the device and not the commands - every command the suite "
              "issues exists here. It is that there is no traffic to learn from, which is "
              "the source-MAC item below.")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED

# ── doc corrections ─────────────────────────────────────────────────────
h(doc, "What \"green\" means here - and what it does not", size=11, space_before=10)
p = doc.add_paragraph()
r = p.add_run("We would rather you get this from us than find it yourselves.")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED
table(doc,
      ["Suite", "Result", "Warnings", "EVPN-behaviour assertions"],
      [["TC01 bring-up", "OK (1 test)", "28", "0"],
       ["TC02 Type-2 MAC/IP", "OK (1 test)", "50", "1"],
       ["TC03 Type-3 IMET", "OK (1 test)", "24", "0"]])
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
r = p.add_run("129 of the 131 reported passes are your framework's own infrastructure "
              "checks - disk space, commit succeeded, IXIA connected, no watchdog reboot, "
              "alarm history clean. The single EVPN assertion (\"no new Type-2 after a "
              "local AC2 to AC3 move\") is the right check, but today it compares an empty "
              "route table with an empty route table, so it cannot fail for the reason it "
              "exists.")
r.font.size = Pt(10.5)
p = doc.add_paragraph()
r = p.add_run("So: the pipeline produces code your device accepts and executes - the "
              "bring-up, the .crt, the .cfg, every CLI command. It does not yet produce "
              "code that verifies EVPN behaviour. Two things stand between the two, one "
              "yours and one ours: there is no traffic to learn from without source-MAC "
              "control on the IXIA, and ate capture's output is not yet fed back into the "
              "generated expectations. Full detail in "
              "02_evidence/evidence_what_the_suites_assert.txt.")
r.font.size = Pt(10.5)

h(doc, "Corrections your EVPN CLI documentation may want")
p = doc.add_paragraph()
r = p.add_run("Found by running against LAB 22, not by reading.")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED
table(doc,
      ["The documents say", "8.7.0 LAB 22 does"],
      [["A vlan-based EVI binds the AC interface",
        "It rejects a port outright: the AC must be a sub-interface "
        "(x-eth 0/0/8.100, l2-transport enable)"],
       ["l2-services evpn <name> has control-word, host "
        "mac-address-duplicate-detection, Advertise-mac, unknow-mac-flooding, "
        "es-waiting-time",
        "The node offers five children only: auto-discovery, interface, "
        "mac-aging-time, mac-limit, service-type"],
       ["interface agg-eth <n> ethernet-segment / lacp-key / lacp-system-mac",
        "No ethernet-segment node under an interface at all - the EVPN "
        "multi-homing configuration is absent from this build"],
       ["service-type accepts vlan-aware-bundle / vlan-bundle",
        "Only port-based and vlan-based"],
       ["mac-limit default 250000",
        "Range <1-250000>, default 65520 - 250000 is the configurable maximum"],
       ["af-l2vpn evpn under BGP",
        "Only under a neighbour or neighbour-group, and only in vrf default"],
       ["show evpn global", "Does not exist — show evpn summary / show evpn detail"],
       ["show evpn bum routing-table",
        "Does not exist — show evpn broadcast-domains carries the BUM label"],
       ["show evpn mac-address-table (hyphen)\nclear evpn mac address-table (space)",
        "Both correct — the product genuinely uses a hyphen for show and a space for clear"],
       ["import-rt / export-rt under the EVI", "They live under auto-discovery"],
       ["show interface … detail", "No detail under show interface"],
       ["show bgp l2vpn evpn table evi evi-name <name>",
        "\"Incomplete path\" until that EVI has BGP EVPN entries; the bare "
        "table evi [detail] form works"]])

# ── asks ────────────────────────────────────────────────────────────────
h(doc, "What we need from you")
bullets(doc, [
    ("A ticket ID — ", "so the branch lands as AUT-nnn / EM-nnnn rather than our "
     "provisional name."),
    ("Source-MAC control on the IXIA — ", "we build the three traffic items in code "
     "over TCL, so no .ixncfg is needed for traffic itself. But ixia_lib.tcl has "
     "editTrafficRawDestMacAddr and no source equivalent, and EVPN learns from source "
     "MACs — so the local MAC-move case (AC2 and AC3 emitting the same source MAC) "
     "cannot be expressed. A src-MAC proc, or an .ixncfg with the three items, "
     "unblocks TC02 and TC03."),
    ("A BGP EVPN peer for this DUT — ", "the four \"show bgp l2vpn evpn table evi "
     "detail\" expectations have nothing to show until a peer exists."),
    ("Confirmation on the absent EVI knobs — ", "control-word, host "
     "mac-address-duplicate-detection, Advertise-mac, unknow-mac-flooding and the "
     "interface ethernet-segment tree are in the CLI doc and not in LAB 22. Either "
     "the document is ahead of the build or the build is missing them; we report it "
     "rather than guess."),
])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run("Closed since the last hand-over: lab-workspace files are no longer "
              "needed - the bring-up completes on pc-3080.")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED

h(doc, "Two things for your attention, neither ours to change", size=11, space_before=10)
bullets(doc, [
    ("exa-il01-ec-3021 has a standing Critical alarm — ", "PSU PSU-1 is Failed. "
     "Pre-existing, and CmpTestCase's @After alarm check will make any suite on that "
     "rig look flaky. pc-3080 is clean: its raised-alarm history is empty."),
    ("cmp/tests/multiCast/MultiCastParams.java — ", "carries a stray "
     "\"import com.sun.javafx.collections.MappingChange;\", an unused IDE auto-import "
     "that only compiled because Oracle JDK 8 shipped JavaFX internals. It fails on "
     "any modern JDK. Left alone rather than shipping an infra fix inside an EVPN branch."),
])

doc.add_page_break()

# ── package ─────────────────────────────────────────────────────────────
h(doc, "The package", space_before=0)
table(doc,
      ["Folder", "Contents"],
      [["01_generated_suite/", "The 8 generated files, in cmp-tests-project layout"],
       ["02_evidence/", "One file per claim above — read these before the code"],
       ["03_test_plan/", "The test plan the code was generated from"],
       ["04_results/", "Full test report (open the .html in a browser)"],
       ["05_git/", "git bundle — 3 commits on auto_develop..ate-m2-evpn-generated-suite"]])

h(doc, "Importing the branch")
p = doc.add_paragraph()
r = p.add_run("git fetch /path/to/evpn-suite.bundle "
              "ate-m2-evpn-generated-suite:<your-branch-name>")
r.font.name = "Consolas"
r.font.size = Pt(9.5)
p = doc.add_paragraph()
r = p.add_run("The branch name in the bundle is provisional. Rename it on import, "
              "or send us a ticket ID and we will.")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED

# ── report ──────────────────────────────────────────────────────────────
h(doc, "Reading the test report")
p = doc.add_paragraph()
r = p.add_run("389 checks — 358 pass, 5 fail, 26 skip.")
r.bold = True
r.font.size = Pt(10.5)
p = doc.add_paragraph()
r = p.add_run("All five failures are code-coverage thresholds (70%) on the CLI wiring "
              "and the device-facing modules, whose network paths are exercised against "
              "real hardware rather than in unit tests. verify.py is the lowest of them "
              "because this round added the session-recovery code that made the "
              "configuration sweep trustworthy. There are no functional test failures "
              "and no lint issues. We are reporting them rather than adjusting the "
              "threshold to hide them.")
r.font.size = Pt(10.5)

h(doc, "The command sweep — now trustworthy in both halves", size=11)
p = doc.add_paragraph()
r = p.add_run("The previous hand-over shipped a sweep of all 123 command templates and "
              "asked you not to act on its configuration-mode half, because it reported "
              "commands missing that the device demonstrably offers. That is fixed and the "
              "cause is understood: a \"?\" on a leaf does not list and return - the CLI "
              "opens an interactive prompt for the value, no prompt character follows, and "
              "the answer was left in the channel for the next probe to collect. Every "
              "later verdict then described the wrong command.")
r.font.size = Pt(10.5)
p = doc.add_paragraph()
r = p.add_run("The sweep now recognises that state, escapes it with Ctrl-C without ever "
              "answering it (this stage is read-only), and proves the channel is back at "
              "its prompt after every probe - this run needed zero recoveries. Twenty "
              "verdicts spanning both halves were then established by hand at the CLI and "
              "compared: 20 of 20 agree. Result: 48 supported, 67 missing, 8 unknown "
              "(02_evidence/evidence_command_verification.txt).")
r.font.size = Pt(10.5)
p = doc.add_paragraph()
r = p.add_run("\"Missing\" means this build does not offer the command - not that your "
              "documentation is wrong. The largest block is the EVPN multi-homing "
              "configuration, which LAB 22 does not expose at all.")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
r = p.add_run("Ilan Ganor · CodeValue · ilan@kamacode.com")
r.font.size = Pt(9.5)
r.font.color.rgb = MUTED

doc.save(OUT)
print("wrote", OUT)
